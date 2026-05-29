"""
포트폴리오 md 파일을 docx로 변환하는 스크립트.
통일된 디자인 시스템 적용:
- 폰트: 본문 Pretendard/맑은 고딕, 코드 Consolas
- 색상: 헤더 #1a1a2e, 본문 #2d2d2d, 서브텍스트 #555555
- 코드 블록: 배경 #f5f5f5, 테두리 #e0e0e0
- 인용문: 좌측 보더 느낌 + 회색 배경
- 테이블: 헤더 #1a1a2e 배경 + 흰색 텍스트
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── 디자인 토큰 ───────────────────────────────────────────────
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x2E)       # 딥 네이비
COLOR_BODY = RGBColor(0x2D, 0x2D, 0x2D)          # 다크 그레이
COLOR_SUB = RGBColor(0x55, 0x55, 0x55)            # 서브 텍스트
COLOR_ACCENT = RGBColor(0x3B, 0x82, 0xF6)        # 블루 액센트
COLOR_QUOTE = RGBColor(0x4B, 0x55, 0x63)         # 인용문 텍스트
COLOR_CODE_TEXT = RGBColor(0x1F, 0x29, 0x37)     # 코드 텍스트
COLOR_TABLE_HEADER_BG = "1A1A2E"                  # 테이블 헤더 배경
COLOR_TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_ALT_BG = "F8FAFC"                     # 테이블 줄무늬
COLOR_CODE_BG = "F5F5F5"                          # 코드 블록 배경
COLOR_QUOTE_BG = "F0F4F8"                         # 인용문 배경

FONT_BODY = "맑은 고딕"
FONT_CODE = "Consolas"

SIZE_H1 = Pt(20)
SIZE_H2 = Pt(14)
SIZE_H3 = Pt(12)
SIZE_BODY = Pt(10)
SIZE_CODE = Pt(8.5)
SIZE_SMALL = Pt(9)
SIZE_TABLE = Pt(9)

SPACING_AFTER_H1 = Pt(12)
SPACING_AFTER_H2 = Pt(8)
SPACING_AFTER_H3 = Pt(6)
SPACING_AFTER_PARA = Pt(6)
SPACING_BEFORE_H2 = Pt(18)
SPACING_BEFORE_H3 = Pt(12)


def _set_cell_shading(cell, color_hex: str):
    """셀 배경색 설정"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_paragraph_shading(paragraph, color_hex: str):
    """단락 배경색 설정"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    paragraph._p.get_or_add_pPr().append(shading_elm)


def _add_horizontal_line(doc):
    """구분선 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="E0E0E0"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def create_docx_from_md(md_path: str, docx_path: str):
    """마크다운 파일을 읽어서 스타일 통일된 docx로 변환"""
    md_content = Path(md_path).read_text(encoding="utf-8")

    doc = Document()

    # ─── 기본 스타일 설정 ───
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = SIZE_BODY
    style.font.color.rgb = COLOR_BODY
    style.paragraph_format.space_after = SPACING_AFTER_PARA
    style.paragraph_format.line_spacing = 1.4

    # 페이지 여백
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = md_content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # ─── 코드 블록 ───
        if line.strip().startswith("```"):
            if in_code_block:
                # 코드 블록 끝 → 렌더링
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    _add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ─── 테이블 ───
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # ─── 빈 줄 ───
        if not line.strip():
            i += 1
            continue

        # ─── 구분선 ───
        if line.strip() == "---":
            _add_horizontal_line(doc)
            i += 1
            continue

        # ─── 헤더 ───
        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue

        # ─── 인용문 ───
        if line.startswith("> "):
            _add_quote(doc, line[2:].strip())
            i += 1
            continue

        # ─── 리스트 ───
        if re.match(r"^[\-\*•]\s", line.strip()):
            text = re.sub(r"^[\-\*•]\s", "", line.strip())
            _add_list_item(doc, text, ordered=False)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s", "", line.strip())
            _add_list_item(doc, text, ordered=True)
            i += 1
            continue

        # ─── 일반 텍스트 ───
        text = line.strip()
        if text:
            _add_paragraph(doc, text)
        i += 1

    # 마지막 테이블 처리
    if in_table:
        _add_table(doc, table_rows)

    doc.save(docx_path)
    print(f"✓ 생성 완료: {docx_path}")


def _add_heading(doc, text: str, level: int):
    """스타일 통일된 헤더"""
    p = doc.add_paragraph()

    if level == 1:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = SPACING_AFTER_H1
        run = p.add_run(_clean_md(text))
        run.font.size = SIZE_H1
        run.font.color.rgb = COLOR_HEADING
        run.bold = True
    elif level == 2:
        p.paragraph_format.space_before = SPACING_BEFORE_H2
        p.paragraph_format.space_after = SPACING_AFTER_H2
        run = p.add_run(_clean_md(text))
        run.font.size = SIZE_H2
        run.font.color.rgb = COLOR_HEADING
        run.bold = True
    elif level == 3:
        p.paragraph_format.space_before = SPACING_BEFORE_H3
        p.paragraph_format.space_after = SPACING_AFTER_H3
        run = p.add_run(_clean_md(text))
        run.font.size = SIZE_H3
        run.font.color.rgb = COLOR_HEADING
        run.bold = True


def _add_paragraph(doc, text: str):
    """본문 단락 (인라인 서식 지원)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = SPACING_AFTER_PARA
    _add_formatted_runs(p, text)


def _add_formatted_runs(paragraph, text: str):
    """인라인 마크다운 서식을 run으로 분리"""
    # 패턴: **bold**, *italic*, `code`, [link](url)
    pattern = r"(\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*|\[(.+?)\]\(.+?\))"
    last_end = 0

    for match in re.finditer(pattern, text):
        # 매치 전 일반 텍스트
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            run = paragraph.add_run(plain)
            run.font.name = FONT_BODY
            run.font.size = SIZE_BODY
            run.font.color.rgb = COLOR_BODY

        full = match.group(0)
        if full.startswith("**"):
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.name = FONT_BODY
            run.font.size = SIZE_BODY
            run.font.color.rgb = COLOR_BODY
        elif full.startswith("`"):
            run = paragraph.add_run(match.group(3))
            run.font.name = FONT_CODE
            run.font.size = SIZE_CODE
            run.font.color.rgb = COLOR_CODE_TEXT
        elif full.startswith("*"):
            run = paragraph.add_run(match.group(4))
            run.italic = True
            run.font.name = FONT_BODY
            run.font.size = SIZE_BODY
            run.font.color.rgb = COLOR_SUB
        elif full.startswith("["):
            run = paragraph.add_run(match.group(5))
            run.font.name = FONT_BODY
            run.font.size = SIZE_BODY
            run.font.color.rgb = COLOR_ACCENT
            run.underline = True

        last_end = match.end()

    # 나머지 텍스트
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = FONT_BODY
        run.font.size = SIZE_BODY
        run.font.color.rgb = COLOR_BODY


def _add_quote(doc, text: str):
    """인용문 — 배경색 + 좌측 인덴트"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    _set_paragraph_shading(p, COLOR_QUOTE_BG)
    run = p.add_run(f"  {_clean_md(text)}")
    run.italic = True
    run.font.size = SIZE_BODY
    run.font.color.rgb = COLOR_QUOTE
    run.font.name = FONT_BODY


def _add_code_block(doc, lines: list):
    """코드 블록 — 배경색 + 모노스페이스"""
    code_text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    _set_paragraph_shading(p, COLOR_CODE_BG)

    run = p.add_run(code_text)
    run.font.name = FONT_CODE
    run.font.size = SIZE_CODE
    run.font.color.rgb = COLOR_CODE_TEXT


def _add_list_item(doc, text: str, ordered: bool):
    """리스트 아이템"""
    style_name = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(3)
    _add_formatted_runs(p, text)


def _add_table(doc, rows: list):
    """테이블 — 헤더 배경 + 줄무늬"""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 테이블 너비 자동
    table.autofit = True

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                continue
            cell = table.cell(row_idx, col_idx)
            # 기존 텍스트 제거
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            clean_text = _clean_md(cell_text)
            run = p.add_run(clean_text)
            run.font.name = FONT_BODY
            run.font.size = SIZE_TABLE

            if row_idx == 0:
                # 헤더 행
                run.bold = True
                run.font.color.rgb = COLOR_TABLE_HEADER_TEXT
                _set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
            else:
                run.font.color.rgb = COLOR_BODY
                # 짝수 행 줄무늬
                if row_idx % 2 == 0:
                    _set_cell_shading(cell, COLOR_TABLE_ALT_BG)

    # 테이블 후 간격
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


def _clean_md(text: str) -> str:
    """마크다운 인라인 서식 제거 (plain text용)"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    return text


if __name__ == "__main__":
    base = Path(__file__).parent

    create_docx_from_md(
        md_path=str(base / "README.md"),
        docx_path=str(base / "QA-TC-Generator-Agent-Portfolio.docx"),
    )

    create_docx_from_md(
        md_path=str(base / "webapp-portfolio.md"),
        docx_path=str(base / "QA-TC-Generator-Webapp-Portfolio.docx"),
    )
